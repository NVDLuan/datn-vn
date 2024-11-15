# from django.http import HttpResponse
# from django.template.loader import render_to_string
# from django.contrib import messages
# import os
# from docx import Document
# from docx.shared import Cm, Pt
# from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# def print_to_docx(modeladmin, request, queryset):
#     """
#     Hàm in dữ liệu ra file docx với đường dẫn template động.
#     """
#     if queryset.count() != 1:
#         messages.error(request, 'Vui lòng chọn một bản ghi duy nhất để in.')
#         return

#     obj = queryset.first()

#     # Lạo đường dẫn động dựa trên tên app, tên model và 'print_template.docx'
#     app_label = obj._meta.app_label
#     model_name = obj._meta.model_name
#     template_path = os.path.join('admin', app_label, model_name, 'template.docx')
#     print(template_path)

    ## CÁCH 1
    # # Tạo tài liệu mới
    # from docx import Document
    # from docx.shared import Cm, Pt, Inches
    # from docx.enum.text import WD_ALIGN_PARAGRAPH
    # from docx.oxml.ns import qn
    # from docx.oxml import OxmlElement
    # from PIL import Image

    # # Tạo tài liệu mới
    # document = Document()
    # sections = document.sections
    # for section in sections:
    #     section.top_margin = Cm(2.5)
    #     section.bottom_margin = Cm(2.5)
    #     section.left_margin = Cm(3)
    #     section.right_margin = Cm(3)

    #     # Thêm viền cho section (khung viền)
    #     sectPr = section._sectPr
    #     pgBorders = OxmlElement('w:pgBorders')
    #     pgBorders.set(qn('w:offsetFrom'), 'page')
    #     for border_name in ['top', 'left', 'bottom', 'right']:
    #         border = OxmlElement(f'w:{border_name}')  # Loại bỏ dấu cách thừa
    #         border.set(qn('w:val'), 'single')
    #         border.set(qn('w:sz'), '8')
    #         border.set(qn('w:space'), '24')
    #         border.set(qn('w:color'), 'auto')
    #         pgBorders.append(border)
    #     sectPr.append(pgBorders)

    # # Tiêu đề
    # heading = document.add_paragraph()
    # heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # run = heading.add_run('BÁO CÁO TIỂU LUẬN')
    # font = run.font
    # font.name = 'Times New Roman'
    # font.size = Pt(20)
    # font.bold = True

    # # Chèn hình ảnh sách mở
    # paragraph = document.add_paragraph()
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # run = paragraph.add_run()
    # image_path = 'sach_mo.jpg'  # Thay bằng đường dẫn đến ảnh của bạn
    # try:
    #     with Image.open(image_path) as img:
    #         width, height = img.size
    #         aspect_ratio = height / width
    #         max_width = Inches(3.5)  # Giới hạn chiều rộng tối đa
    #         new_width = min(max_width, Inches(width / 96))  # Chuyển đổi pixel sang inch (96 DPI)
    #         new_height = new_width * aspect_ratio
    #         run.add_picture(image_path, width=new_width, height=new_height)
    # except Exception as e:
    #     print(f"Lỗi khi chèn hình ảnh: {e}")

    # # Thông tin bên dưới
    # info_paragraph = document.add_paragraph()
    # info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # info_paragraph.add_run(f'Giáo viên hướng dẫn:\t{obj.ma_co_quan}\n')
    # info_paragraph.add_run('Sinh viên thực hiện:\t....................................................\n')
    # info_paragraph.add_run('Lớp:\t....................................................\n')
    # info_paragraph.add_run('Nhóm:\t....................................................\n')

    # # Chèn logo
    # paragraph = document.add_paragraph()
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # run = paragraph.add_run()
    # logo_path = 'logo.png'  # Thay bằng đường dẫn đến logo của bạn
    # try:
    #     with Image.open(logo_path) as img:
    #         width, height = img.size
    #         aspect_ratio = height / width
    #         max_height = Inches(0.8)  # Giới hạn chiều cao tối đa
    #         new_height = min(max_height, Inches(height / 96))
    #         new_width = new_height / aspect_ratio
    #         run.add_picture(logo_path, width=new_width, height=new_height)
    # except Exception as e:
    #     print(f"Lỗi khi chèn logo: {e}")

    # # Năm học
    # year_paragraph = document.add_paragraph()
    # year_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # run = year_paragraph.add_run('Năm học 2020-2021')
    # font = run.font
    # font.name = 'Times New Roman'
    # font.size = Pt(12)

    # # CÁCH 2
    # import re
    # from docx import Document

    # def docx_replace_regex(doc_obj, regex , replace):

    #     for p in doc_obj.paragraphs:
    #         if regex.search(p.text):
    #             inline = p.runs
    #             # Loop added to work with runs (strings with same style)
    #             for i in range(len(inline)):
    #                 if regex.search(inline[i].text):
    #                     text = regex.sub(replace, inline[i].text)
    #                     inline[i].text = text

    #     for table in doc_obj.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 docx_replace_regex(cell, regex , replace)



    # regex1 = re.compile(r"your regex")
    # replace1 = r"your replace string"
    # filename = "test.docx"
    # document = Document("/home/tuna/Desktop/records-management-mini/home/templates/"+template_path)
    # docx_replace_regex(document, regex1 , replace1)
   
    # # Tạo response để download file docx
    # response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    # response['Content-Disposition'] = f'attachment; filename="{obj}-{obj.pk}.docx"'  # Đặt tên file
    # document.save(response)
    # return response

# print_to_docx.short_description = "In thành file docx"


from django.http import HttpResponse
from django.template.loader import render_to_string
from django.contrib import messages
import os
from docx import Document
from django.conf import settings
from django.shortcuts import get_object_or_404
from django.apps import apps
from .models import (
    DonVi, ThoiHanLuuTru, LoaiVanBan, PhanLoaiTaiLieuLuuTru,
    KhoLuuTru, GiaLuuTru, PhongLuuTru, MucLucHoSo, HoSo,
    VanBan, TapTinDinhKem, UserDonVi
)

def print_to_docx(request, app_label, model_name, pk):
    """
    Hàm in dữ liệu ra file docx với đường dẫn template động.
    """
    # Lấy model class từ app_label và model_name
    model = apps.get_model(app_label, model_name)

    # Lấy object dựa trên pk
    obj = get_object_or_404(model, pk=pk)

    # Tạo đường dẫn động dựa trên tên app, tên model và 'print_template.docx'
    template_path = os.path.join('admin', app_label, model_name, 'template.docx')
    full_template_path = os.path.join(settings.BASE_DIR, 'home','templates', template_path)

    # Kiểm tra và tạo thư mục nếu chưa tồn tại
    os.makedirs(os.path.dirname(full_template_path), exist_ok=True)

    # Kiểm tra xem template có tồn tại hay không
    if not os.path.exists(full_template_path):
        messages.error(request, f"Không tìm thấy template tại: {full_template_path}")
        return HttpResponse("Không tìm thấy template")

    import re
    from docx import Document

    def docx_replace_regex_new(doc, replacements_dict):
        from docx import Document
        from lxml import etree

        for regex_pattern, replacement in replacements_dict.items():
            # Load the document

            # Get the root element of the document
            root = doc._element

            # Example: Print out the XML as a string
            print(etree.tostring(root, pretty_print=True).decode())

            # Example: Find all paragraphs and replace text
            for p in root.findall('.//w:p', namespaces=root.nsmap):
                for t in p.findall('.//w:t', namespaces=root.nsmap):
                    if regex_pattern in t.text:
                        t.text = t.text.replace(str(regex_pattern), str(replacement))

            # # Save the modified document
            # doc.save(r"C:\Users\nguye\OneDrive\template.docx")

    def docx_replace_regex(doc_obj, replacements_dict):
        """
        Replaces text in a docx document based on multiple regex patterns and their corresponding replacements.

        Args:
            doc_obj: A Document object from the python-docx library.
            replacements_dict: A dictionary where keys are regex patterns (strings) and values are their replacements.
        """

        for regex_pattern, replacement in replacements_dict.items():
            regex = re.compile(regex_pattern)

            for p in doc_obj.paragraphs:
                if regex.search(p.text):
                    for run in p.runs:  
                        if regex.search(run.text):
                            run.text = regex.sub(str(replacement), str(run.text))

            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        docx_replace_regex(cell, replacements_dict)


    # Định nghĩa regex và replace cho từng model
    if isinstance(obj, DonVi):
        dic_regex = {
            r"{{ ten }}": obj.ten,
            r"{{ ma }}": obj.ma,
            r"{{ dia_chi }}": obj.dia_chi,
            r"{{ dien_thoai }}": obj.dien_thoai,
            r"{{ email }}": obj.email,
            r"{{ fax }}": obj.fax,
            r"{{ ghi_chu }}": obj.ghi_chu,
        }
        
        # ... thêm các regex và replace khác cho DonVi nếu cần
    elif isinstance(obj, ThoiHanLuuTru):
        dic_regex = {
            r"{{ ten }}": obj.ten,
            r"{{ so_nam }}": obj.so_nam,
            r"{{ ghi_chu }}": obj.ghi_chu,
        }
        # ... thêm các regex và replace khác cho ThoiHanLuuTru nếu cần
    # ... (thêm elif cho các model khác)
    elif isinstance(obj, HoSo):
        dic_regex = {
            r"{{ tieu_de }}": obj.tieu_de,
            r"{{ ho_so_so }}": obj.ho_so_so,
            r"yyy": obj.phong_luu_tru,
            r"xxx": obj.phong_luu_tru.kho_luu_tru,
            r"zzz": obj.hop_luu_tru.hop_so,
            r"{{ hop_luu_tru.tu_nam }}": obj.hop_luu_tru.tu_nam,
            r"{{ ghi_chu }}": obj.ghi_chu,
        }
        # ... thêm các regex và replace khác cho ThoiHanLuuTru nếu cần
    # ... (thêm elif cho các model khác)

    document = Document(full_template_path)
    print(document)
    docx_replace_regex_new(document, dic_regex)
   
    # Tạo response để download file docx
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{obj}-{obj.pk}.docx"'  # Đặt tên file
    document.save(response)
    return response
