from django.http import HttpResponse
from django.template.loader import render_to_string
from django.contrib import messages
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from import_export.admin import ImportExportModelAdmin

from .models import *
from .resources import *

from django.contrib import admin
from django.db.models import Q
from django.contrib import admin
from django.contrib.auth.admin import UserAdmin as BaseUserAdmin
from django.contrib.auth.models import User as DjangoUser
from django.contrib.auth.models import Group
from rangefilter.filter import DateRangeFilter, DateTimeRangeFilter

from .models import (
    DonVi, ThoiHanLuuTru, LoaiVanBan, PhanLoaiTaiLieuLuuTru,
    KhoLuuTru, GiaLuuTru, PhongLuuTru, MucLucHoSo, HoSo,
    VanBan, TapTinDinhKem, UserDonVi, HopLuuTru
)
from django.conf import settings
from django.core.cache import cache

def print_to_docx(modeladmin, request, queryset):
    """
    Hàm in dữ liệu ra file docx với đường dẫn template động.
    """
    if queryset.count() != 1:
        messages.error(request, 'Vui lòng chọn một bản ghi duy nhất để in.')
        return

    obj = queryset.first()

    # Lạo đường dẫn động dựa trên tên app, tên model và 'print_template.docx'
    app_label = obj._meta.app_label
    model_name = obj._meta.model_name
    template_path = os.path.join('admin', app_label, model_name, 'template.docx')
    full_template_path = os.path.join(settings.BASE_DIR, 'home','templates', template_path)

    # Kiểm tra và tạo thư mục nếu chưa tồn tại
    os.makedirs(os.path.dirname(full_template_path), exist_ok=True)

    # Kiểm tra xem template có tồn tại hay không
    if not os.path.exists(full_template_path):
        messages.error(request, f"Không tìm thấy template tại: {full_template_path}")
        return

    import re
    from docx import Document

    def docx_replace_regex(doc_obj, regex, replace):
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace)

    # Định nghĩa regex và replace cho từng model
    if isinstance(obj, DonVi):
        regex1 = re.compile(r"{{ ten_don_vi }}")
        replace1 = obj.ten_don_vi
        # ... thêm các regex và replace khác cho DonVi nếu cần
    elif isinstance(obj, ThoiHanLuuTru):
        regex1 = re.compile(r"{{ ten }}")
        replace1 = obj.ten
        # ... thêm các regex và replace khác cho ThoiHanLuuTru nếu cần
    # ... (thêm elif cho các model khác)

    document = Document(full_template_path)
    docx_replace_regex(document, regex1, replace1)
   
    # Tạo response để download file docx
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{obj}-{obj.pk}.docx"'  # Đặt tên file
    document.save(response)
    return response

print_to_docx.short_description = "In thành file docx"

# ModelAdmin cho các model
class DonViAdmin(ImportExportModelAdmin):
    resource_class = DonViResource
    filter_horizontal = ('groups',) 
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'trang_thai',
    )
    search_fields = ['ten_don_vi', 'ma']  # Thêm trường tìm kiếm

    # Cache kết quả queryset trong 1 giờ
    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = f'donvi_queryset_{request.user.id}'
        if request.user.is_superuser:
            cached_qs = cache.get(cache_key)
            if cached_qs is None:
                cached_qs = qs
                cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
            return cached_qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        filtered_qs = qs.filter(phongluutru__don_vi__in=don_vi_ids)
        cache.set(cache_key, filtered_qs, 60 * 60)
        return filtered_qs

    @admin.display(description='Số lượng phòng')
    def so_luong_phong(self, obj):
        cache_key = f'donvi_{obj.id}_so_luong_phong'
        so_luong = cache.get(cache_key)
        if so_luong is None:
            so_luong = obj.phongluutru_set.count()
            cache.set(cache_key, so_luong, 60 * 60)  # Cache số lượng phòng trong 1 giờ
        return so_luong

class ThoiHanLuuTruAdmin(ImportExportModelAdmin):
    resource_class = ThoiHanLuuTruResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'trang_thai',
    )
    search_fields = ['ten']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'thoihanluutru_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        return cached_qs

class LoaiVanBanAdmin(ImportExportModelAdmin):
    resource_class = LoaiVanBanResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'trang_thai',
    )
    search_fields = ['ten', 'ten_viet_tat']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'loaivanban_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        return cached_qs

class PhanLoaiTaiLieuLuuTruAdmin(ImportExportModelAdmin):
    resource_class = PhanLoaiTaiLieuLuuTruResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'trang_thai',
    )
    search_fields = ['ten']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'phanloaitailieu_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        return cached_qs

class KhoLuuTruAdmin(ImportExportModelAdmin):
    resource_class = KhoLuuTruResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'trang_thai',
    )
    search_fields = ['ten_kho']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'kholuutru_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        return cached_qs

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.user.is_superuser:
            return qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return qs.filter(phongluutru__don_vi__in=don_vi_ids)

class GiaLuuTruAdmin(ImportExportModelAdmin):
    resource_class = GiaLuuTruResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'kho_luu_tru',
        'trang_thai',
    )
    search_fields = ['ten', 'gia_so']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'gialuutru_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        return cached_qs

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.user.is_superuser:
            return qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return qs.filter(kho_luu_tru__phongluutru__don_vi__in=don_vi_ids)

class PhongLuuTruAdmin(ImportExportModelAdmin):
    resource_class = PhongLuuTruResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'don_vi', 'kho_luu_tru', 
        'trang_thai',
        'nam_bat_dau', 'nam_ket_thuc', # Filter theo năm
    )
    search_fields = ['name', 'phong_so']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'phongluutru_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        if request.user.is_superuser:
            return cached_qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return cached_qs.filter(don_vi__in=don_vi_ids)

    @admin.display(description='Số lượng hồ sơ')
    def so_luong_ho_so(self, obj):
        cache_key = f'phongluutru_{obj.id}_so_luong_ho_so'
        so_luong = cache.get(cache_key)
        if so_luong is None:
            so_luong = obj.hoso_set.count()
            cache.set(cache_key, so_luong, 60 * 60)  # Cache số lượng hồ sơ trong 1 giờ
        return so_luong

class MucLucHoSoAdmin(ImportExportModelAdmin):
    resource_class = MucLucHoSoResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'phong_luu_tru', 
        'trang_thai',
    )
    search_fields = ['muc_luc_so']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'mucluchoso_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        if request.user.is_superuser:
            return cached_qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return cached_qs.filter(phong_luu_tru__don_vi__in=don_vi_ids)

class HoSoAdmin(ImportExportModelAdmin):
    resource_class = HoSoResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'phong_luu_tru', 'hop_luu_tru', 'muc_luc_ho_so', 'phan_loai_tai_lieu_luu_tru',
        'thoi_han_bao_quan', 
        'trang_thai',
        ('thoi_gian_bat_dau', DateRangeFilter), # Filter theo khoảng thời gian
        ('thoi_gian_ket_thuc', DateRangeFilter),
    )
    search_fields = ['tieu_de', 'ho_so_so', 'ma_ho_so', 'tu_khoa']  # Thêm trường tìm kiếm

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'hoso_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        if request.user.is_superuser:
            return cached_qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return cached_qs.filter(phong_luu_tru__don_vi__in=don_vi_ids)

class VanBanAdmin(ImportExportModelAdmin):
    resource_class = VanBanResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'ho_so', 'loai_van_ban', 'thoi_han_bao_quan', 
        ('ngay_ban_hanh', DateRangeFilter),
    )
    search_fields = ['so_van_ban', 'ky_hieu_cua_van_ban', 'trich_yeu_noi_dung', 'tu_khoa']

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'vanban_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        return cached_qs

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.user.is_superuser:
            return qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return qs.filter(ho_so__phong_luu_tru__don_vi__in=don_vi_ids)

class UserDonViInline(admin.TabularInline):
    model = UserDonVi
    extra = 1

class UserAdmin(BaseUserAdmin):
    inlines = (UserDonViInline,)

class HopLuuTruAdmin(ImportExportModelAdmin):
    resource_class = HopLuuTruResource
    actions = [print_to_docx]
    list_filter = (
        ('created_at', DateTimeRangeFilter),
        ('updated_at', DateTimeRangeFilter),
        'phong_luu_tru', 'kho_luu_tru', 'gia_luu_tru', 'muc_luc_ho_so',
        'trang_thai',
        ('tu_nam', DateRangeFilter),
        ('den_nam', DateRangeFilter)
    )
    search_fields = ['hop_so', 'mo_ta']

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        cache_key = 'hopluutru_queryset'
        cached_qs = cache.get(cache_key)
        if cached_qs is None:
            cached_qs = qs
            cache.set(cache_key, cached_qs, 60 * 60)  # Cache queryset trong 1 giờ
        if request.user.is_superuser:
            return cached_qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return cached_qs.filter(phong_luu_tru__don_vi__in=don_vi_ids)


class TapTinDinhKemAdmin(admin.ModelAdmin):
    list_display = ('van_ban', 'tep_tin')
    search_fields = ['van_ban__so_van_ban', 'van_ban__ky_hieu_cua_van_ban']
    readonly_fields = ('van_ban',)

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.user.is_superuser:
            return qs
        don_vi_ids = request.user.userdonvi_set.values_list('don_vi', flat=True)
        return qs.filter(van_ban__ho_so__phong_luu_tru__don_vi__in=don_vi_ids)

# Đăng ký ModelAdmin
admin.site.unregister(DjangoUser)
admin.site.register(DjangoUser, UserAdmin)
admin.site.register(DonVi, DonViAdmin)
admin.site.register(ThoiHanLuuTru, ThoiHanLuuTruAdmin)
admin.site.register(LoaiVanBan, LoaiVanBanAdmin)
admin.site.register(PhanLoaiTaiLieuLuuTru, PhanLoaiTaiLieuLuuTruAdmin)
admin.site.register(KhoLuuTru, KhoLuuTruAdmin)
admin.site.register(GiaLuuTru, GiaLuuTruAdmin)
admin.site.register(PhongLuuTru, PhongLuuTruAdmin)
admin.site.register(MucLucHoSo, MucLucHoSoAdmin)
admin.site.register(HoSo, HoSoAdmin)
admin.site.register(VanBan, VanBanAdmin)
admin.site.register(HopLuuTru, HopLuuTruAdmin)
admin.site.register(TapTinDinhKem, TapTinDinhKemAdmin)
# admin.site.register(VanBan, VanBanAdmin)